"""
parser.py
=========
Resume text extraction for PDF, DOCX, and TXT formats.

Supported formats:
    - PDF  → PyPDF2
    - DOCX → python-docx
    - TXT  → built-in

Install:
    pip install PyPDF2 python-docx
"""

import os
import re


def extract_text_from_pdf(filepath: str) -> str:
    """
    Extract text from PDF resume.
    Priority: pdfplumber → PyPDF2 → pdfminer
    pdfplumber preserves newlines best (important for name detection).
    """
    # ── Method 1: pdfplumber (best newline preservation) ──────
    try:
        import pdfplumber
        text  = ""
        pages = 0
        with pdfplumber.open(filepath) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        if text.strip():
            print(f"✓  PDF extracted: {len(text)} chars ({pages} pages)")
            return text
    except ImportError:
        pass
    except Exception as e:
        print(f"   ⚠  pdfplumber error: {e}")

    # ── Method 2: PyPDF2 ──────────────────────────────────────
    try:
        import PyPDF2
        text = ""
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        if text.strip():
            print(f"✓  PDF extracted: {len(text)} chars ({len(reader.pages)} pages)")
            return text
    except ImportError:
        pass
    except Exception as e:
        print(f"   ⚠  PyPDF2 error: {e}")

    # ── Method 3: pdfminer ────────────────────────────────────
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(filepath)
        if text and text.strip():
            print(f"✓  PDF extracted: {len(text)} chars (pdfminer)")
            return text
    except ImportError:
        pass
    except Exception as e:
        print(f"   ⚠  pdfminer error: {e}")

    print("✗  No PDF library available. Run: pip install pdfplumber")
    return ""


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX resume."""
    try:
        from docx import Document
        doc  = Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        print(f"✓  DOCX extracted: {len(text)} chars ({len(doc.paragraphs)} paragraphs)")
        return text
    except ImportError:
        print("✗  python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        print(f"✗  DOCX extraction error: {e}")
        return ""


def extract_text_from_txt(filepath: str) -> str:
    """Extract text from a plain TXT resume."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
        print(f"✓  TXT extracted: {len(text)} chars")
        return text
    except Exception as e:
        print(f"✗  TXT extraction error: {e}")
        return ""


def extract_text(filepath: str) -> str:
    """
    Auto-detect file format and extract text from resume.

    Args:
        filepath (str): Path to resume file (PDF / DOCX / TXT)

    Returns:
        str: Extracted raw text
    """
    if not os.path.exists(filepath):
        print(f"✗  File not found: {filepath}")
        return ""

    ext = os.path.splitext(filepath)[1].lower()

    print(f"\n📄  Extracting text from: {os.path.basename(filepath)}")

    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        return extract_text_from_txt(filepath)
    else:
        print(f"✗  Unsupported format: {ext}. Use PDF, DOCX, or TXT.")
        return ""